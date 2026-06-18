import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()

    # Titre
    title = doc.add_heading("Formation Complète : Tests d'API avec Postman (Projet E-Santé)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Ce document est un cours détaillé, étape par étape, pour vous apprendre à configurer et maîtriser les tests d'API sur votre application E-Santé à l'aide de Postman.")

    # Section 1
    doc.add_heading("Chapitre 1 : Les Concepts Fondamentaux", level=1)
    p = doc.add_paragraph()
    p.add_run("1. Qu'est-ce que Postman ? ").bold = True
    p.add_run("Postman est un outil graphique permettant aux développeurs de requêter des APIs (Web Services), sans avoir besoin de coder le Front-End (Flutter React etc.).\n")
    p.add_run("2. Méthodes HTTP Principales :\n").bold = True
    p.add_run("   • GET : Récupérer des informations (ex: lister les hôpitaux).\n")
    p.add_run("   • POST : Créer de l'information (ex: prendre un rendez-vous, se connecter).\n")
    p.add_run("   • PUT / PATCH : Modifier des informations existantes.\n")
    p.add_run("   • DELETE : Supprimer ou désactiver des données.")

    # Section 1.5 - Setup
    doc.add_heading("Chapitre 2 : Configuration de votre Espace Postman", level=1)
    doc.add_paragraph("L'une des plus grandes forces de Postman est l'utilisation des Variables d'Environnement.")
    doc.add_paragraph("1. En haut à droite dans Postman, cliquez sur l'icône de l'œil (Environment Quick Look) et faites 'Add'.\n2. Nommez votre environnement : 'E-Santé Local'.\n3. Créez les variables suivantes :")
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Variable'
    hdr_cells[1].text = 'Valeur Initiale'
    
    row1 = table.add_row().cells
    row1[0].text = 'base_url'
    row1[1].text = 'http://127.0.0.1:8000/api'
    
    row2 = table.add_row().cells
    row2[0].text = 'token'
    row2[1].text = '(Laissez vide pour l\'instant)'

    doc.add_paragraph("\nMaintenant, sélectionnez 'E-Santé Local' dans le menu déroulant d'environnement en haut à droite.")

    # Section 2 - Authentification
    doc.add_heading("Chapitre 3 : Obtenir et utiliser un Token JWT (L'Authentification)", level=1)
    doc.add_paragraph("L'API d'E-Santé est sécurisée par le système JWT (JSON Web Tokens). Pour accéder aux routes protégées, vous devez vous identifier.", style='Intense Quote')
    
    doc.add_heading("Étape 3.1 : Obtenir le Token", level=2)
    doc.add_paragraph("• Méthode : POST\n• URL : {{base_url}}/token/")
    doc.add_paragraph("Allez dans l'onglet 'Body' > 'raw' > 'JSON' et insérez :")
    
    code1 = doc.add_paragraph()
    code1.add_run('{\n  "email": "votre_email@patient.com",\n  "password": "votre_mot_de_passe"\n}').font.name = 'Courier New'
    
    doc.add_paragraph("Testez ! La réponse contiendra un \"access\". C'est votre token.")

    doc.add_heading("Étape 3.2 : Automatiser la sauvegarde du Token (Astuce Pro) ⭐", level=2)
    doc.add_paragraph("Au lieu de copier/coller le token à chaque fois, allez dans l'onglet 'Tests' de votre requête /token/ et ajoutez ce script JavaScript :")
    
    code2 = doc.add_paragraph()
    code2.add_run('var jsonData = pm.response.json();\nif (jsonData.access) {\n    pm.environment.set("token", jsonData.access);\n    console.log("Token enregistré !");\n}').font.name = 'Courier New'
    
    doc.add_paragraph("Ainsi, à chaque connexion, Postman sauvegarde le jeton pour vos autres requêtes !")

    # Section 3 - Tester les ressources E-Santé
    doc.add_heading("Chapitre 4 : Cas Pratiques avec les fonctionnalités de l'Application", level=1)
    
    # 4.1 Prox
    doc.add_heading("Scénario A : Recherche des Hôpitaux à Proximité", level=2)
    doc.add_paragraph("• Méthode : GET\n• URL : {{base_url}}/hopitaux/nearby/?lat=6.37&lng=2.39&radius=10\n• Authorization : Allez dans l'onglet 'Authorization', choisissez 'Bearer Token' et mettez `{{token}}` dans la case.")
    doc.add_paragraph("Ce que ça teste : Vérifie que le backend liste bien les hôpitaux de notre Database les plus proches, et valide que SEUL un Patient peut y accéder.")

    # 4.2 RDV
    doc.add_heading("Scénario B : Prise de Rendez-vous (Création complexe)", level=2)
    doc.add_paragraph("• Méthode : POST\n• URL : {{base_url}}/rendezvous/\n• Authorization : Bearer {{token}}")
    doc.add_paragraph("Onglet Body (JSON) :")
    code3 = doc.add_paragraph()
    code3.add_run('{\n  "hopital": 1,\n  "service": 2,\n  "medecin": 3,\n  "date_prevue": "2026-05-15",\n  "heure_debut": "09:00:00",\n  "motif": "Consultation générale"\n}').font.name = 'Courier New'
    doc.add_paragraph("Testez la validation : Changez l'heure à 02:00:00. L'API devrait vous rejeter car le médecin n'est pas disponible la nuit !")

    # 4.3 BioTrack
    doc.add_heading("Scénario C : Importer des résultats d'analyses", level=2)
    doc.add_paragraph("• Méthode : POST\n• URL : {{base_url}}/resultats/\n• Rôle Requis : Laborantin")
    doc.add_paragraph("Comme ce endpoint utilise des fichiers PDF, allez dans l'onglet 'Body', choisissez 'form-data' :")
    doc.add_paragraph("   - Clé 'fichier' -> Changez le type 'Text' en 'File' -> Sélectionnez un PDF.\n   - Clé 'patient_id' -> Identifiant du patient (ex: 2).\n   - Clé 'type_analyse' -> Ex: 'Bilan Sanguin'.")

    # Conclusion
    doc.add_heading("Chapitre 5 : Intégrer des vrais tests automatiques", level=1)
    doc.add_paragraph("Dans l'onglet 'Tests' de Postman, vous pouvez écrire du code pour qu'en un clic, Postman valide que tout va bien (utile avant d'envoyer l'app en production). Exemple :")
    code4 = doc.add_paragraph()
    code4.add_run('// Vérifie qu\'on a bien un statut 200 OK\npm.test("Status code is 200", function () {\n    pm.response.to.have.status(200);\n});\n\n// Vérifie que le temps de réponse est rapide (ex: distance géographique)\npm.test("Response time is under 500ms", function () {\n    pm.expect(pm.response.responseTime).to.be.below(500);\n});').font.name = 'Courier New'

    doc.add_paragraph("\n\nFélicitations ! Avec ces scénarios, vous pouvez simuler tout le cycle E-Santé : Le médecin, le patient, le laboratoire et l'administrateur sans toucher au code Flutter.")

    output_path = os.path.abspath(r"E:\Soutenance\Dossiers\Local\Backend\Cours_Postman_ESante.docx")
    doc.save(output_path)
    print(f"Le fichier WORD a été généré avec succès dans : {output_path}")

if __name__ == "__main__":
    main()
